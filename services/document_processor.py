"""
Document Processor Service

Handles text extraction from PDF and DOCX documents, with OCR fallback for scanned PDFs.
"""

import io
from pathlib import Path
from typing import Optional, Tuple
import re

# PDF Processing
from PyPDF2 import PdfReader

# DOCX Processing
from docx import Document

# OCR Processing
try:
    import pytesseract
    from PIL import Image
    from pdf2image import convert_from_path, convert_from_bytes
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


class DocumentProcessor:
    """
    Processes RFP documents (PDF, DOCX) and extracts text.
    
    Features:
    - PDF text extraction with PyPDF2
    - DOCX text extraction with python-docx
    - OCR fallback for scanned PDFs using Tesseract
    - Auto-detection of scanned documents
    """
    
    # Minimum chars per page to consider it text-based (not scanned)
    MIN_CHARS_PER_PAGE = 100
    
    def __init__(self, tesseract_path: Optional[str] = None):
        """
        Initialize the document processor.
        
        Args:
            tesseract_path: Optional path to Tesseract executable
        """
        if tesseract_path and OCR_AVAILABLE:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
    
    def process_file(self, file_path: Path) -> dict:
        """
        Process a document file and extract text.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dict with extracted text and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        suffix = file_path.suffix.lower()
        
        if suffix == ".pdf":
            return self._process_pdf(file_path)
        elif suffix in [".docx", ".doc"]:
            return self._process_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
    
    def process_bytes(self, file_bytes: bytes, filename: str) -> dict:
        """
        Process document from bytes.
        
        Args:
            file_bytes: Document content as bytes
            filename: Original filename (for format detection)
            
        Returns:
            Dict with extracted text and metadata
        """
        suffix = Path(filename).suffix.lower()
        
        if suffix == ".pdf":
            return self._process_pdf_bytes(file_bytes)
        elif suffix in [".docx", ".doc"]:
            return self._process_docx_bytes(file_bytes)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
    
    def _process_pdf(self, file_path: Path) -> dict:
        """Process a PDF file."""
        with open(file_path, "rb") as f:
            return self._process_pdf_bytes(f.read(), str(file_path))
    
    def _process_pdf_bytes(self, pdf_bytes: bytes, source: str = "bytes") -> dict:
        """Process PDF from bytes."""
        result = {
            "source": source,
            "format": "pdf",
            "text": "",
            "pages": [],
            "page_count": 0,
            "ocr_used": False,
            "extraction_method": "native",
            "warnings": []
        }
        
        try:
            reader = PdfReader(io.BytesIO(pdf_bytes))
            result["page_count"] = len(reader.pages)
            
            all_text = []
            low_text_pages = 0
            
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                page_text = self._clean_text(page_text)
                
                result["pages"].append({
                    "page_number": i + 1,
                    "text": page_text,
                    "char_count": len(page_text)
                })
                
                all_text.append(page_text)
                
                # Check if this might be a scanned page
                if len(page_text) < self.MIN_CHARS_PER_PAGE:
                    low_text_pages += 1
            
            result["text"] = "\n\n".join(all_text)
            
            # If most pages have low text, try OCR
            if low_text_pages > result["page_count"] / 2:
                if OCR_AVAILABLE:
                    result["warnings"].append("Low text extraction detected. Attempting OCR...")
                    ocr_result = self._ocr_pdf_bytes(pdf_bytes)
                    if ocr_result and len(ocr_result) > len(result["text"]):
                        result["text"] = ocr_result
                        result["ocr_used"] = True
                        result["extraction_method"] = "ocr"
                else:
                    result["warnings"].append(
                        "Low text extraction detected. OCR not available - install pytesseract and pdf2image."
                    )
        
        except Exception as e:
            result["warnings"].append(f"PDF processing error: {str(e)}")
            # Try OCR as fallback
            if OCR_AVAILABLE:
                try:
                    result["text"] = self._ocr_pdf_bytes(pdf_bytes)
                    result["ocr_used"] = True
                    result["extraction_method"] = "ocr_fallback"
                except Exception as ocr_e:
                    result["warnings"].append(f"OCR fallback also failed: {str(ocr_e)}")
        
        return result
    
    def _ocr_pdf_bytes(self, pdf_bytes: bytes) -> str:
        """Extract text from PDF using OCR."""
        if not OCR_AVAILABLE:
            raise RuntimeError("OCR not available")
        
        try:
            # Convert PDF to images
            images = convert_from_bytes(pdf_bytes, dpi=300)
            
            all_text = []
            for i, image in enumerate(images):
                # Run OCR on each page
                text = pytesseract.image_to_string(image)
                text = self._clean_text(text)
                all_text.append(text)
            
            return "\n\n".join(all_text)
        
        except Exception as e:
            raise RuntimeError(f"OCR processing failed: {str(e)}")
    
    def _process_docx(self, file_path: Path) -> dict:
        """Process a DOCX file."""
        with open(file_path, "rb") as f:
            return self._process_docx_bytes(f.read(), str(file_path))
    
    def _process_docx_bytes(self, docx_bytes: bytes, source: str = "bytes") -> dict:
        """Process DOCX from bytes."""
        result = {
            "source": source,
            "format": "docx",
            "text": "",
            "sections": [],
            "paragraph_count": 0,
            "ocr_used": False,
            "extraction_method": "native",
            "warnings": []
        }
        
        try:
            doc = Document(io.BytesIO(docx_bytes))
            
            paragraphs = []
            current_section = {"heading": None, "content": []}
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Check if this is a heading
                if para.style.name.startswith("Heading"):
                    if current_section["content"]:
                        result["sections"].append(current_section)
                    current_section = {"heading": text, "content": []}
                else:
                    current_section["content"].append(text)
                    paragraphs.append(text)
            
            # Add last section
            if current_section["content"]:
                result["sections"].append(current_section)
            
            result["text"] = "\n\n".join(paragraphs)
            result["paragraph_count"] = len(paragraphs)
            
            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
            
            if table_text:
                result["text"] += "\n\n[Table Content]\n" + "\n".join(table_text)
        
        except Exception as e:
            result["warnings"].append(f"DOCX processing error: {str(e)}")
        
        return result
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text."""
        if not text:
            return ""
        
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Fix common OCR issues
        text = text.replace('|', 'l')  # Common OCR mistake
        
        # Remove excessive newlines but preserve paragraph breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()


# Module-level instance for convenience
_processor = None

def get_processor() -> DocumentProcessor:
    """Get or create the document processor instance."""
    global _processor
    if _processor is None:
        _processor = DocumentProcessor()
    return _processor


def extract_text(file_path: Path) -> str:
    """
    Convenience function to extract text from a document.
    
    Args:
        file_path: Path to the document
        
    Returns:
        Extracted text
    """
    processor = get_processor()
    result = processor.process_file(file_path)
    return result["text"]


def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """
    Convenience function to extract text from document bytes.
    
    Args:
        file_bytes: Document content as bytes
        filename: Original filename
        
    Returns:
        Extracted text
    """
    processor = get_processor()
    result = processor.process_bytes(file_bytes, filename)
    return result["text"]
